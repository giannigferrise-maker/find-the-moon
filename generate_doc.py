"""
Generate find-the-moon-design.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins (1 inch all around) ────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0a, 0x16, 0x28)   # dark night-sky blue
INDIGO    = RGBColor(0x3b, 0x0d, 0xa8)   # deep indigo accent
PURPLE    = RGBColor(0x6d, 0x28, 0xd9)   # vibrant purple
SLATE     = RGBColor(0x47, 0x55, 0x69)   # body text grey
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GOLD      = RGBColor(0xca, 0x8a, 0x04)
LIGHT_BG  = RGBColor(0xf0, 0xf4, 0xff)   # subtle lavender tint for shaded rows

# ── Helper: shade a table cell ───────────────────────────────────────────────
def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

# ── Helper: set paragraph background (page-level shading via frame) ──────────
def set_para_shading(para, hex_colour):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    pPr.append(shd)

# ── Helper: add a styled heading ─────────────────────────────────────────────
def add_heading(text, level=1):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(17)
        run.font.color.rgb = RGBColor(0x3b, 0x0d, 0xa8)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '6d28d9')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = PURPLE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
    return p

# ── Helper: body paragraph ───────────────────────────────────────────────────
def add_body(text, indent=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(11)
    run.font.color.rgb = SLATE
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p

# ── Helper: bullet point ─────────────────────────────────────────────────────
def add_bullet(label, detail, level=0):
    p    = doc.add_paragraph(style='List Bullet')
    r_lbl = p.add_run(label)
    r_lbl.bold            = True
    r_lbl.font.color.rgb  = NAVY
    r_lbl.font.size       = Pt(11)
    if detail:
        r_sep = p.add_run(' — ')
        r_sep.font.color.rgb = SLATE
        r_sep.font.size      = Pt(11)
        r_det = p.add_run(detail)
        r_det.font.color.rgb = SLATE
        r_det.font.size      = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    if level:
        p.paragraph_format.left_indent = Inches(0.3 * level)

# ── Helper: numbered step ────────────────────────────────────────────────────
def add_step(number, text):
    p   = doc.add_paragraph(style='List Number')
    num = p.add_run(f'Step {number}: ')
    num.bold           = True
    num.font.color.rgb = PURPLE
    num.font.size      = Pt(11)
    body = p.add_run(text)
    body.font.color.rgb = SLATE
    body.font.size      = Pt(11)
    p.paragraph_format.space_after = Pt(4)

# ── Helper: callout box (indented, lightly shaded paragraph) ─────────────────
def add_callout(text):
    p = doc.add_paragraph()
    set_para_shading(p, 'EEF2FF')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run('💡  ' + text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)
    run.font.italic    = True

# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run('🌙')
r.font.size = Pt(60)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run('Find the Moon')
t.bold           = True
t.font.size      = Pt(36)
t.font.color.rgb = RGBColor(0x3b, 0x0d, 0xa8)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run('Project Design Document')
s.font.size      = Pt(16)
s.font.color.rgb = PURPLE

doc.add_paragraph()  # spacer

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tl = tagline_p.add_run(
    'A single-page web application that tells you exactly\n'
    'where to look in the sky to find the moon — right now.'
)
tl.font.size      = Pt(13)
tl.font.color.rgb = SLATE
tl.font.italic    = True

for _ in range(5):
    doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = meta_p.add_run(
    f'Version 1.0   ·   {datetime.date.today().strftime("%B %d, %Y")}   ·   Single-File Web App'
)
meta.font.size      = Pt(10)
meta.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════
add_heading('1. Project Overview', level=1)

add_body(
    'Find the Moon is a simple, beautiful web application that answers one magical question: '
    '"Where is the moon right now?" In just a few seconds, the app figures out your location, '
    'does the astronomical math behind the scenes, and shows you exactly which direction to look '
    'and how high up in the sky the moon will be. No telescopes, no star charts, no complicated '
    'apps — just open it in your browser and go outside.'
)
add_body(
    'The app was designed to feel wonder-inspiring, especially for children. Its visual style '
    'automatically shifts between a glittering night sky (when it\'s dark outside) and a soft '
    'daytime sky (when the sun is up), making every visit feel immersive and alive.'
)

add_heading('Who Is It For?', level=2)
add_bullet('Families and children', 'A fun way to go moon-hunting together on a clear evening.')
add_bullet('Casual stargazers', 'Quick answer for anyone curious about the moon\'s location without learning astronomy.')
add_bullet('Educators', 'A great classroom or homework tool for discussing moon phases, the solar system, and navigation.')
add_bullet('Photographers', 'Quickly scout the moon\'s position before a shoot.')
add_bullet('Anyone with curiosity', 'Because the night sky is endlessly fascinating.')

add_heading('Platform & Accessibility', level=2)
add_body(
    'Find the Moon runs entirely in any modern web browser — Chrome, Safari, Firefox, Edge — '
    'on phones, tablets, and desktop computers alike. No app store download is needed. '
    'No account, no subscription, no tracking. Because it is a single self-contained HTML file, '
    'it can even work offline once downloaded.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2. FEATURES
# ════════════════════════════════════════════════════════════════════
add_heading('2. Features', level=1)

add_heading('2.1  Location Detection (GPS)', level=2)
add_body(
    'When you first open the app, it offers to detect your location automatically using your '
    'device\'s built-in GPS. You simply tap "Use My Location" and — if you grant permission — '
    'the app instantly knows where you are in the world. No typing required. This works on '
    'phones, tablets, and most laptops that have a location sensor.'
)
add_callout(
    'Your location data never leaves your device. The app calculates everything locally in your '
    'browser — it is never sent to any server or stored anywhere.'
)

add_heading('2.2  Zip Code Lookup', level=2)
add_body(
    'If you prefer not to share your GPS location, or if you want to check the moon from a '
    'different city, you can simply type any US zip code into the search box. The app looks up '
    'the geographic coordinates for that zip code and uses them for the moon calculations. '
    'This is great for planning ahead — for example, checking tonight\'s moon position for '
    'a camping trip destination.'
)

add_heading('2.3  Moon Direction Compass', level=2)
add_body(
    'The heart of the app is the moon direction compass. It shows a beautiful circular compass '
    'with a moon icon needle pointing in the exact direction you need to look. Below it, the '
    'app spells out the direction in plain English — for example, "Look Southeast" — along with '
    'the precise compass bearing in degrees. You don\'t need to know anything about degrees; '
    'the plain-language direction is all you need.'
)

add_heading('2.4  Live Compass (Mobile Devices)', level=2)
add_body(
    'On smartphones and tablets that have a built-in compass sensor (magnetometer), the app '
    'offers an interactive live compass experience. When enabled, a real compass rose appears '
    'that rotates in real time as you physically move your phone. A small moon icon sits on '
    'the compass ring at the moon\'s exact position. A fixed purple arrow at the top of the '
    'compass represents "the direction you are currently pointing your phone." Simply rotate '
    'your body until the moon icon lines up with that arrow — and you\'re facing the moon!'
)
add_bullet('iOS devices', 'The app requests motion sensor permission, as required by Apple\'s privacy guidelines.', level=0)
add_bullet('Android devices', 'Activates automatically when you tap the button — no extra permission step needed.', level=0)

add_heading('2.5  Moon Altitude Display', level=2)
add_body(
    'Knowing which direction to look is only half the story — you also need to know how far '
    'up to tilt your head. The app shows the moon\'s current altitude (height above the '
    'horizon) as a number in degrees. It also draws a semicircular arc — like a protractor '
    'laid on its flat side — with a glowing moon dot showing exactly how high up the moon '
    'currently sits in the sky arc from horizon to directly overhead.'
)
add_body(
    'For example, 0° means the moon is right at the horizon (just rising or setting), '
    '45° means it\'s halfway up the sky, and 90° means it\'s directly overhead.'
)

add_heading('2.6  Moon Phase Display', level=2)
add_body(
    'The app shows the current moon phase with a hand-drawn graphical representation of the '
    'moon\'s illuminated shape — from a thin crescent sliver all the way to a glowing full moon. '
    'The phase name is shown in plain English (e.g., "Waxing Crescent," "Full Moon," '
    '"Waning Gibbous") along with the percentage of the moon\'s face that is currently lit up '
    'by the sun.'
)
add_bullet('Waxing', 'The moon is growing — more of it becomes lit each night.')
add_bullet('Waning', 'The moon is shrinking — less of it is lit each night.')
add_bullet('Crescent', 'A thin sliver is visible.')
add_bullet('Quarter', 'Half of the moon\'s face is lit.')
add_bullet('Gibbous', 'More than half is lit, but not yet full.')

add_heading('2.7  Visibility Status', level=2)
add_body(
    'The app clearly tells you whether the moon is currently above the horizon (visible) or '
    'below it (not yet visible). If the moon is visible, it shows when it will set. '
    'If the moon is below the horizon, it tells you exactly what time the moon will rise '
    'so you know when to head outside.'
)

add_heading('2.8  Automatic Day / Night Theme', level=2)
add_body(
    'The app automatically detects whether it is currently daytime or nighttime at your '
    'location and switches its visual appearance accordingly.'
)
add_bullet('Night theme', 'A rich, deep navy sky filled with hundreds of softly twinkling stars, rendered as tiny points of light across the background.')
add_bullet('Day theme', 'A bright blue sky gradient with gently drifting white clouds that drift slowly across the screen.')
add_body(
    'The theme is determined by whether the sun is currently above or below the horizon at '
    'your location — not just by the clock — so it is accurate even near sunrise and sunset.'
)

add_heading('2.9  Auto-Refresh', level=2)
add_body(
    'The moon moves continuously across the sky. To keep the information accurate without '
    'you needing to do anything, the app automatically recalculates and updates all of the '
    'moon\'s data every 60 seconds in the background.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  3. HOW IT WORKS
# ════════════════════════════════════════════════════════════════════
add_heading('3. How It Works', level=1)

add_heading('3.1  Getting Your Location', level=2)
add_body(
    'The app needs to know two things about your location: your latitude (how far north or '
    'south you are) and your longitude (how far east or west you are). It gets these in '
    'one of two ways:'
)
add_step(1, 'GPS — Your browser asks your device\'s location sensor for your current coordinates. '
            'This is instant and very precise.')
add_step(2, 'Zip Code — You type in a US zip code, and the app contacts a free public service '
            'called Zippopotam.us to look up the latitude and longitude of that zip code\'s '
            'town centre. This takes about one second.')

add_heading('3.2  How the Moon\'s Position Is Calculated', level=2)
add_body(
    'Once the app knows where you are on Earth, it uses a well-tested astronomy library called '
    'SunCalc.js to figure out exactly where the moon is in the sky. Here is what happens, '
    'in plain English:'
)
add_step(1, 'The app notes the exact current date and time down to the second.')
add_step(2, 'It passes your location and the current time to SunCalc.js.')
add_step(3, 'SunCalc.js uses established astronomical formulas to calculate the moon\'s '
            'position relative to your specific spot on Earth. '
            'These are the same kinds of formulas used by observatories.')
add_step(4, 'The result is the moon\'s azimuth (the compass direction, measured in degrees '
            'clockwise from North) and its altitude (how high above the horizon it is, '
            'in degrees).')
add_step(5, 'These two numbers are turned into the visual compass needle, the altitude arc, '
            'and the plain-English direction text.')

add_callout(
    'Azimuth and altitude are the two coordinates of the "horizontal coordinate system" — '
    'the simplest way to describe where something is in the sky as seen from a specific '
    'point on Earth. Think of azimuth as the direction on a compass, and altitude as the '
    'angle you tilt your head up.'
)

add_heading('3.3  Moon Phase Calculation', level=2)
add_body(
    'SunCalc.js also calculates the moon\'s illumination — in other words, what fraction of '
    'the moon\'s face is currently lit up by the sun. This is determined by the angle between '
    'the Earth, the Moon, and the Sun at any given moment. From this angle, the app determines:'
)
add_bullet('The illumination fraction', '0% = new moon (dark), 100% = full moon (completely lit).')
add_bullet('Whether the moon is waxing or waning', 'Based on which half of the lunar cycle we\'re in.')
add_bullet('The phase name', 'Such as "Waxing Crescent" or "Waning Gibbous."')
add_body(
    'The graphical moon shape is then drawn as an SVG (a scalable graphic) directly in the '
    'browser, using the mathematical shape of the illuminated crescent or gibbous.'
)

add_heading('3.4  Rise and Set Times', level=2)
add_body(
    'SunCalc.js also calculates what time the moon rises and sets for your location on any '
    'given day. If the moon is currently below the horizon, the app finds the next scheduled '
    'rise time — including checking the following day if necessary — and displays it.'
)

add_heading('3.5  The Live Compass', level=2)
add_body(
    'On mobile devices, the compass works by reading the phone\'s built-in orientation sensor '
    '(called a magnetometer or compass sensor). The browser provides this data through a '
    'standard feature called the DeviceOrientation API:'
)
add_bullet('iPhone / iPad (iOS)', 'Reports compass heading directly in degrees from North via a property called webkitCompassHeading.')
add_bullet('Android phones', 'Reports orientation data through the deviceorientationabsolute event, which the app converts to a compass heading.')
add_body(
    'The app reads the compass heading many times per second and redraws the compass rose '
    'accordingly, so it moves smoothly in real time as you rotate your phone. The moon icon '
    'on the compass ring stays fixed at the moon\'s true azimuth, while the compass rose '
    'rotates beneath it — exactly like a real physical compass.'
)

add_heading('3.6  Zip Code to Coordinates', level=2)
add_body(
    'When a zip code is entered, the app makes a simple request to the free Zippopotam.us '
    'API — a public database of US (and international) postal codes with their geographic '
    'coordinates. The response comes back in about half a second and includes the town name, '
    'state, latitude, and longitude. No API key is required, and the service is completely free.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  4. LIBRARIES & DEPENDENCIES
# ════════════════════════════════════════════════════════════════════
add_heading('4. Libraries & Dependencies', level=1)

add_body(
    'Find the Moon is intentionally lightweight. It uses only one external JavaScript library '
    'and one external web API, both of which are completely free and require no account or '
    'API key to use.'
)

# Table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.autofit = True

hdr_cells = tbl.rows[0].cells
headers = ['Name', 'Type', 'What It Does', 'Cost / License']
for i, h in enumerate(headers):
    shade_cell(hdr_cells[i], '3B0DA8')
    run = hdr_cells[i].paragraphs[0].add_run(h)
    run.bold           = True
    run.font.color.rgb = WHITE
    run.font.size      = Pt(11)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ('SunCalc.js',
     'JavaScript Library',
     'Calculates the position of the moon (and sun) for any location and time. '
     'Provides azimuth, altitude, illumination fraction, phase angle, and rise/set times.',
     'Free & Open Source\n(BSD-2-Clause)'),
    ('Zippopotam.us API',
     'Web API',
     'Converts a US zip code into geographic coordinates (latitude & longitude) '
     'plus a human-readable city and state name.',
     'Free, no key needed\nPublic service'),
    ('Browser APIs\n(built-in)',
     'Browser Feature',
     'Geolocation API (GPS), DeviceOrientation API (live compass sensor), '
     'Canvas 2D API (drawing graphics), and RequestAnimationFrame (smooth animation). '
     'All built into every modern browser — nothing to install.',
     'Free\nPart of web standards'),
]

for i, (name, type_, does, cost) in enumerate(rows_data):
    row = tbl.add_row().cells
    fill = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        shade_cell(cell, fill)
    row[0].paragraphs[0].add_run(name).font.size = Pt(10.5)
    row[0].paragraphs[0].runs[0].bold = True
    row[1].paragraphs[0].add_run(type_).font.size = Pt(10.5)
    row[2].paragraphs[0].add_run(does).font.size = Pt(10.5)
    row[3].paragraphs[0].add_run(cost).font.size = Pt(10.5)
    row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x15, 0x80, 0x3d)

doc.add_paragraph()  # spacer

add_callout(
    'Because SunCalc.js is loaded from a public content-delivery network (CDN), '
    'the app requires an internet connection to load for the first time. '
    'Once loaded, all moon calculations happen entirely in the browser with no '
    'further network requests (except for zip code lookups).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  5. FILE STRUCTURE
# ════════════════════════════════════════════════════════════════════
add_heading('5. File Structure', level=1)

add_body(
    'One of the design goals for Find the Moon was radical simplicity. The entire application '
    'is a single file. There is no build process, no package manager, no server required.'
)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(['File', 'Size', 'Purpose']):
    shade_cell(hdr2[i], '3B0DA8')
    r = hdr2[i].paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)
    hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

files_data = [
    ('find-the-moon/\n  index.html', '~31 KB', 'The entire application. Contains all HTML structure, CSS styling, '
     'and JavaScript logic in a single self-contained file. Open this file in '
     'any browser to run the app.'),
    ('find-the-moon/\n  find-the-moon-design.docx', '~generated', 'This design document (the file you are reading right now).'),
]
for i, (file, size, purpose) in enumerate(files_data):
    row = tbl2.add_row().cells
    fill = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        shade_cell(cell, fill)
    r0 = row[0].paragraphs[0].add_run(file)
    r0.font.size = Pt(10); r0.font.name = 'Courier New'; r0.bold = True
    row[1].paragraphs[0].add_run(size).font.size = Pt(10.5)
    row[2].paragraphs[0].add_run(purpose).font.size = Pt(10.5)

doc.add_paragraph()

add_heading('Inside index.html', level=2)
add_body('The single HTML file is organised into three clear sections:')
add_bullet('HTML (structure)',
    'Defines the layout of the page — the header, location input fields, compass card, '
    'phase display, altitude arc, visibility card, and the live compass. '
    'Approximately 80 lines.')
add_bullet('CSS (styling)',
    'All visual design: colours, fonts, card layouts, animations (stars, clouds, needle spin), '
    'day and night theme variables, responsive layout for mobile screens. '
    'Approximately 380 lines.')
add_bullet('JavaScript (logic)',
    'All interactivity and calculations: GPS detection, zip code lookup, SunCalc.js calls, '
    'compass needle animation, altitude arc drawing, moon SVG rendering, live compass, '
    'theme switching, and auto-refresh. Approximately 350 lines.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  6. FUTURE ENHANCEMENT IDEAS
# ════════════════════════════════════════════════════════════════════
add_heading('6. Future Enhancement Ideas', level=1)

add_body(
    'The current version of Find the Moon is intentionally focused and simple. '
    'Here are some directions the project could grow in the future:'
)

add_heading('Astronomy & Sky Features', level=2)
add_bullet('Interactive star map',
    'Show the major constellations around the moon\'s current position, with names '
    'and mythology — turning the app into a light star-chart companion.')
add_bullet('Planet finder',
    'Extend the same direction/altitude display to the visible planets '
    '(Venus, Mars, Jupiter, Saturn) so users can find all the bright objects at once.')
add_bullet('Sun position',
    'Add a "Find the Sun" mode using the same SunCalc.js library, '
    'useful for photographers and solar panel owners.')
add_bullet('Lunar calendar',
    'A small monthly calendar showing the upcoming full moon, new moon, and eclipse dates.')
add_bullet('Augmented reality mode',
    'Use the phone\'s camera feed as a background and overlay the compass and moon '
    'indicator on top of the real world, like a simple AR sky guide.')

add_heading('Location & Personalization', level=2)
add_bullet('International postal codes',
    'The Zippopotam.us API supports many countries. Extending the input to accept '
    'UK postcodes, Canadian postal codes, etc. would make the app globally useful.')
add_bullet('Search by city name',
    'Let users type "Paris" or "Tokyo" instead of a zip code, using a geocoding API.')
add_bullet('Saved locations',
    'Remember favourite locations (home, cabin, grandparents\' house) in the browser\'s '
    'local storage so users can switch between them quickly.')

add_heading('Social & Sharing', level=2)
add_bullet('Share the moon',
    'A "Share" button that generates a link or an image card showing the current '
    'moon phase and direction, ready to post on social media.')
add_bullet('Tonight\'s moon alert',
    'A browser notification (if the user opts in) that pings when the moon rises '
    'or when it reaches its highest point in the sky tonight.')
add_bullet('Moonrise / moonset countdown',
    'A live countdown timer to the next moonrise or moonset event.')

add_heading('Design & Experience', level=2)
add_bullet('Animated moon phase transitions',
    'Smoothly animate the SVG moon shape as the phase changes over time, '
    'making it feel like a living astronomical illustration.')
add_bullet('Sound design',
    'Optional gentle ambient sounds — a light breeze, soft night crickets, or a '
    'quiet musical note when the moon is found — to add to the magical feeling.')
add_bullet('Kid mode',
    'A special extra-simplified view with larger text, a talking guide character, '
    'and fun moon facts for younger children.')
add_bullet('Dark-sky quality indicator',
    'Use the user\'s location to fetch local light-pollution data and tell them '
    'how good their stargazing conditions are tonight.')
add_bullet('Offline / PWA support',
    'Package the app as a Progressive Web App so it can be added to the home screen '
    'and work without an internet connection using cached data.')

add_heading('Technical Improvements', level=2)
add_bullet('Atmospheric refraction correction',
    'The moon appears slightly higher in the sky near the horizon due to how the '
    'atmosphere bends light. Adding this correction would make the altitude reading '
    'even more accurate.')
add_bullet('Compass calibration hint',
    'On Android, compass accuracy varies. Adding an on-screen prompt to "draw a figure 8" '
    '(the standard compass calibration gesture) would improve accuracy.')
add_bullet('Multiple language support',
    'Translate the interface into Spanish, French, German, Japanese, and other languages '
    'to make it accessible to a global audience.')

# ════════════════════════════════════════════════════════════════════
#  FOOTER NOTE
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()

close_p = doc.add_paragraph()
close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
close_p.paragraph_format.space_before = Pt(80)
c = close_p.add_run('🌕')
c.font.size = Pt(40)

close_p2 = doc.add_paragraph()
close_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ct = close_p2.add_run('"The moon is a loyal companion.\nIt never leaves. It\'s always there, watching, steadfast, knowing us in our light and dark moments."')
ct.font.size = Pt(12)
ct.font.italic = True
ct.font.color.rgb = SLATE

close_p3 = doc.add_paragraph()
close_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca = close_p3.add_run('— Tahereh Mafi')
ca.font.size = Pt(10)
ca.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_paragraph()
doc.add_paragraph()

end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ep = end_p.add_run(f'Find the Moon  ·  Design Document  ·  {datetime.date.today().strftime("%Y")}')
ep.font.size = Pt(9)
ep.font.color.rgb = RGBColor(0xc0, 0xca, 0xd8)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save('/Users/gianniferrise/find-the-moon/find-the-moon-design.docx')
print('Document saved successfully.')
